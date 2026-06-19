from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/opciones_pago_bajo_costo_mexico.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.allow_autofit = False

    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.append(tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        tbl_grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[idx]))


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.1):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_run(paragraph, text, bold=False, color=None, size=None):
    run = paragraph.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)
    return run


def add_bullet(doc, text, level=0):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    set_paragraph_spacing(paragraph, after=8, line=1.167)
    paragraph.add_run(text)
    return paragraph


def add_number(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    set_paragraph_spacing(paragraph, after=8, line=1.167)
    paragraph.add_run(text)
    return paragraph


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Opciones de pago bajo costo - México")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    paragraph = cell.paragraphs[0]
    set_paragraph_spacing(paragraph, after=4, line=1.1)
    add_run(paragraph, title, bold=True, color="1F3A5F")
    paragraph.add_run(" " + body)
    doc.add_paragraph()


def add_decision_table(doc):
    doc.add_heading("Comparación rápida", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    widths = [2100, 3630, 3630]
    set_table_geometry(table, widths)
    headers = ["Criterio", "CoDi vía proveedor", "SPEI con CLABE única"]
    for idx, text in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_shading(cell, "F2F4F7")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, text, bold=True, color="1F4D78")
    set_repeat_table_header(table.rows[0])

    rows = [
        (
            "Experiencia",
            "QR o notificación de cobro; el cliente autoriza desde su app bancaria.",
            "El cliente transfiere a una CLABE o referencia única desde su banco.",
        ),
        (
            "Costo",
            "Puede ser muy bajo; Banxico/CoDi no es el costo principal, pero el proveedor puede cobrar.",
            "Suele ser más barato que tarjeta; depende de tarifa fija, mensualidad y proveedor.",
        ),
        (
            "Confirmación",
            "Buena si el proveedor envía evento/webhook de pago aceptado.",
            "Muy buena si cada orden tiene CLABE o referencia única y webhook de conciliación.",
        ),
        (
            "Fricción",
            "Baja en desktop con QR; en mobile requiere push/link para no obligar a escanear la misma pantalla.",
            "Media: copiar datos, abrir banco, transferir y esperar confirmación.",
        ),
        (
            "Mejor caso",
            "Tickets bajos o medios, clientes bancarizados y búsqueda fuerte de menor comisión.",
            "Tickets medios o altos, necesidad de compatibilidad amplia con bancos y conciliación robusta.",
        ),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            paragraph = cells[idx].paragraphs[0]
            set_paragraph_spacing(paragraph, after=0, line=1.1)
            if idx == 0:
                add_run(paragraph, text, bold=True)
            else:
                paragraph.add_run(text)
    doc.add_paragraph()


def build():
    doc = Document()
    style_document(doc)
    add_footer(doc)

    title = doc.add_paragraph()
    title.style = doc.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Opciones de pago bajo costo en México")
    run.font.name = "Calibri Light"
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
    set_paragraph_spacing(title, after=3, line=1.0)

    subtitle = doc.add_paragraph()
    add_run(
        subtitle,
        "Comparación práctica entre CoDi vía proveedor y SPEI con CLABE única para reducir comisiones sin perder confirmación automática.",
        color="555555",
    )
    set_paragraph_spacing(subtitle, after=12, line=1.15)

    add_callout(
        doc,
        "Recomendación corta:",
        "usar transferencia SPEI con CLABE única como opción de bajo costo principal y evaluar CoDi como flujo rápido si el proveedor ofrece QR/push y webhook confiable.",
    )

    doc.add_heading("Contexto", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "El objetivo no es eliminar toda comisión, sino bajar el costo por cobro frente a tarjeta o Mercado Pago sin volver manual la operación. "
        "La clave técnica es que cada pago pueda identificarse contra una orden y que el backend reciba una confirmación automática antes de liberar la compra."
    )

    add_decision_table(doc)

    doc.add_heading("Opción 1: CoDi vía proveedor", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "CoDi funciona sobre la infraestructura de SPEI. En un flujo web, el comercio genera un QR dinámico o una notificación de cobro; el cliente autoriza desde su app bancaria y el sistema confirma el pago."
    )

    doc.add_heading("Ventajas", level=2)
    for item in [
        "Puede tener costo muy bajo porque el pago no pasa por redes de tarjeta.",
        "Permite confirmación automática si el proveedor entrega webhook o evento de pago aceptado.",
        "La acreditación suele ser rápida porque opera sobre rails bancarios.",
        "Reduce riesgo de contracargos de tarjeta, ya que el pago es autorizado como transferencia.",
        "Es atractivo para tickets bajos o medios donde una comisión de 3% a 4% impacta mucho el margen.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("Desventajas", level=2)
    for item in [
        "Menor familiaridad para muchos compradores frente a tarjeta, Mercado Pago u OXXO.",
        "En celular, un QR en la misma pantalla es incómodo; conviene que el proveedor soporte push o link de autorización.",
        "La integración puede requerir contrato empresarial, onboarding, pruebas, requisitos fiscales y tarifas no públicas.",
        "No ofrece meses sin intereses ni beneficios de tarjeta.",
        "Puede tener límites por operación según banco/proveedor; conviene confirmarlo antes de elegirlo para tickets altos.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("Opción 2: SPEI con CLABE única", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "En este modelo, el proveedor genera una CLABE o referencia única para cada orden o cliente. El comprador transfiere desde su banco y el proveedor notifica al backend cuando el dinero entra, permitiendo conciliación automática."
    )

    doc.add_heading("Ventajas", level=2)
    for item in [
        "Más universal que CoDi: cualquier cliente con banca mexicana puede hacer SPEI.",
        "Suele funcionar mejor para tickets medios y altos.",
        "La conciliación puede ser muy robusta si cada orden tiene CLABE o referencia única.",
        "Puede ser mucho más barato que tarjeta, especialmente si el proveedor cobra fijo bajo por transacción.",
        "No depende de que el usuario conozca CoDi; solo necesita saber transferir.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("Desventajas", level=2)
    for item in [
        "Más fricción: copiar CLABE, abrir banco, transferir y volver al sitio.",
        "Hay que manejar casos operativos: monto incorrecto, pago tardío, pago duplicado, pago parcial o pago de más.",
        "El checkout se siente menos instantáneo que tarjeta si el banco demora la confirmación.",
        "El proveedor puede cobrar mensualidad, mínimo mensual, tarifa fija por transacción o tarifa por API.",
        "Los reembolsos se vuelven más manuales que con una pasarela de tarjeta.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("Flujo técnico recomendado", level=1)
    for item in [
        "Crear orden en estado pendiente_pago.",
        "Solicitar al proveedor un QR CoDi o una CLABE/referencia única para esa orden.",
        "Mostrar instrucciones y un tiempo de vencimiento claro en el checkout.",
        "Recibir webhook de pago confirmado en el backend.",
        "Consultar el pago en la API del proveedor antes de marcar la orden como pagada.",
        "Actualizar la orden de forma idempotente para evitar duplicados.",
    ]:
        add_number(doc, item)

    doc.add_heading("Preguntas para proveedores", level=1)
    for item in [
        "¿Cuál es la comisión por pago SPEI/CoDi y cuál es el cargo fijo por transacción?",
        "¿Hay mensualidad, mínimo mensual, costo de alta o costo de API?",
        "¿El sistema genera CLABE única, referencia única, QR dinámico, push o link?",
        "¿El webhook se dispara solo cuando el pago está confirmado o también cuando está pendiente?",
        "¿Cuánto tarda la liquidación y en qué cuenta se deposita?",
        "¿Cómo se manejan pagos por monto incorrecto, pagos vencidos y reembolsos?",
        "¿Qué requisitos fiscales, bancarios y de volumen piden para México?",
    ]:
        add_bullet(doc, item)

    doc.add_heading("Decisión sugerida", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "Si los tickets son bajos o medios y el proveedor ofrece buen flujo push/link, CoDi puede ser una opción muy atractiva. "
        "Si necesitás compatibilidad amplia, tickets más altos y conciliación fuerte, SPEI con CLABE única parece la primera opción a priorizar. "
        "La estrategia más sana es mantener tarjeta/Mercado Pago como alternativa de conversión y ofrecer transferencia bancaria como opción recomendada o con precio preferencial."
    )

    doc.add_heading("Fuentes para revisar", level=1)
    sources = [
        ("Banxico - CoDi", "https://www.banxico.org.mx/sistemas-de-pago/codi-cobro-digital-banco-me.html"),
        ("STP - CoDi Cobro Digital", "https://stp.mx/soluciones/codi-cobro-digital/"),
        ("Fintoc - Transfers Overview", "https://docs.fintoc.com/docs/transfers-overview"),
        ("DiMo", "https://www.dimo.org.mx/"),
    ]
    for name, url in sources:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        add_hyperlink(paragraph, name, url)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build()
