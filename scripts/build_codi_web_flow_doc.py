from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/flujo_implementacion_codi_web.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=120):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.allow_autofit = False

    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.append(tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        tbl_grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[idx]))


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_spacing(paragraph, before=0, after=6, line=1.1):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_run(paragraph, text, bold=False, color=None, size=None):
    run = paragraph.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)
    return run


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    set_spacing(p, after=6, line=1.167)
    p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    set_spacing(p, after=6, line=1.167)
    p.add_run(text)
    return p


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 14, 7),
        ("Heading 2", 13, "2E74B5", 10, 5),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_footer(doc):
    footer = doc.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Flujo CoDi web")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def add_callout(doc, label, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360], indent=0)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    set_spacing(p, after=2, line=1.1)
    add_run(p, label, bold=True, color="1F3A5F")
    p.add_run(" " + body)
    doc.add_paragraph()


def add_flow_diagram(doc):
    doc.add_heading("Flujo visual", level=1)
    steps = [
        ("1", "Checkout", "Cliente elige CoDi en la web."),
        ("2", "Crear orden", "Backend crea orden pendiente y llama al proveedor."),
        ("3", "QR / push", "Proveedor devuelve QR, link o solicitud push."),
        ("4", "Autorizar", "Cliente confirma pago en app bancaria."),
        ("5", "Webhook", "Proveedor notifica pago confirmado."),
        ("6", "Liberar", "Backend valida y marca la orden como pagada."),
    ]
    table = doc.add_table(rows=2, cols=3)
    table.style = "Table Grid"
    set_table_geometry(table, [3120, 3120, 3120], indent=0)
    for idx, (num, title, body) in enumerate(steps):
        row = 0 if idx < 3 else 1
        col = idx if idx < 3 else idx - 3
        cell = table.cell(row, col)
        set_cell_shading(cell, "E8EEF5")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(p, after=3, line=1.05)
        add_run(p, num + ". " + title, bold=True, color="0B2545", size=12)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(p2, after=0, line=1.05)
        add_run(p2, body, color="202020", size=10)
    doc.add_paragraph()


def add_endpoints_table(doc):
    doc.add_heading("Endpoints mínimos", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    set_table_geometry(table, [2050, 2100, 2510, 2700])
    headers = ["Endpoint", "Quién lo llama", "Qué hace", "Resultado"]
    for idx, text in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_shading(cell, "F2F4F7")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, text, bold=True, color="1F4D78")
    set_repeat_table_header(table.rows[0])
    rows = [
        (
            "POST /pay/codi",
            "Frontend",
            "Crea orden o intento de pago CoDi con monto, moneda y order_id.",
            "Devuelve QR, link, expiración e ID del proveedor.",
        ),
        (
            "GET /order/status",
            "Frontend",
            "Consulta estado mientras el cliente está en pantalla de espera.",
            "Devuelve pendiente, pagado, expirado o fallido.",
        ),
        (
            "POST /webhook/codi",
            "Proveedor",
            "Recibe evento de pago. Debe validar firma y consultar API del proveedor.",
            "Actualiza orden de forma idempotente.",
        ),
        (
            "POST /order/cancel",
            "Frontend/Admin",
            "Cancela intento pendiente cuando vence o el cliente cambia método.",
            "Orden queda lista para nuevo intento.",
        ),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            p = cells[idx].paragraphs[0]
            set_spacing(p, after=0, line=1.05)
            if idx == 0:
                add_run(p, text, bold=True, color="0B2545", size=9.5)
            else:
                add_run(p, text, size=9.5)
    doc.add_paragraph()


def add_state_table(doc):
    doc.add_heading("Estados recomendados", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_geometry(table, [2600, 3000, 3760])
    headers = ["Estado", "Cuándo se usa", "Acción del sistema"]
    for idx, text in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_shading(cell, "F2F4F7")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, text, bold=True, color="1F4D78")
    set_repeat_table_header(table.rows[0])
    rows = [
        ("created", "Orden creada, sin intento de pago.", "Permite elegir método."),
        ("pending_payment", "QR/push generado y no pagado aún.", "Muestra espera, QR, countdown y polling suave."),
        ("paid", "Webhook validado y monto correcto.", "Libera producto/servicio y envía confirmación."),
        ("expired", "El QR o intento venció.", "Permite generar un nuevo intento."),
        ("failed", "Proveedor rechazó o no pudo completar.", "Muestra alternativa de pago."),
        ("manual_review", "Monto incorrecto, duplicado o pago fuera de tiempo.", "No liberar automáticamente; revisar operación."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            p = cells[idx].paragraphs[0]
            set_spacing(p, after=0, line=1.05)
            if idx == 0:
                add_run(p, text, bold=True, color="0B2545", size=9.5)
            else:
                add_run(p, text, size=9.5)
    doc.add_paragraph()


def add_sequence_table(doc):
    doc.add_heading("Secuencia técnica", level=1)
    steps = [
        "El cliente llega al checkout y selecciona CoDi.",
        "El frontend llama al backend con el ID de la orden.",
        "El backend valida stock, monto, moneda, usuario y vencimiento.",
        "El backend crea un intento de pago con el proveedor CoDi.",
        "El proveedor responde con QR, push, link o payload de cobro.",
        "La web muestra pantalla de espera con QR, monto exacto y timer.",
        "El cliente autoriza desde su app bancaria.",
        "El proveedor envía webhook de cambio de estado.",
        "El backend valida firma/origen del webhook y consulta el pago por API.",
        "Si el pago está confirmado y el monto coincide, la orden pasa a pagada.",
    ]
    for step in steps:
        add_number(doc, step)


def build():
    doc = Document()
    style_document(doc)
    add_footer(doc)

    title = doc.add_paragraph()
    title.style = doc.styles["Title"]
    run = title.add_run("Implementación web de CoDi vía proveedor")
    run.font.name = "Calibri Light"
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
    set_spacing(title, after=3, line=1.0)

    subtitle = doc.add_paragraph()
    add_run(
        subtitle,
        "Guía de flujo para integrar cobros CoDi en un ecommerce con QR/push, webhook y conciliación automática.",
        color="555555",
    )
    set_spacing(subtitle, after=12, line=1.15)

    add_callout(
        doc,
        "Idea central:",
        "tu web nunca debe marcar una orden como pagada solo porque mostró un QR. La orden se libera únicamente después de recibir un webhook y validar el pago consultando al proveedor.",
    )

    doc.add_heading("Arquitectura mínima", level=1)
    for item in [
        "Frontend de checkout: muestra CoDi como método, QR/push, monto exacto, timer y estado de espera.",
        "Backend: crea la orden, habla con el proveedor, recibe webhooks y decide si la orden queda pagada.",
        "Base de datos: guarda orden, intento de pago, estado, monto, moneda, vencimiento e ID del proveedor.",
        "Proveedor CoDi: genera el cobro, conecta con la app bancaria del cliente y notifica resultado.",
        "Panel/admin: permite revisar pagos duplicados, vencidos, incorrectos o con error.",
    ]:
        add_bullet(doc, item)

    add_flow_diagram(doc)
    add_sequence_table(doc)
    add_endpoints_table(doc)
    add_state_table(doc)

    doc.add_heading("Datos que deberías guardar", level=1)
    for item in [
        "order_id interno.",
        "payment_attempt_id interno.",
        "provider_payment_id o transaction_id.",
        "monto exacto y moneda MXN.",
        "estado de orden y estado del intento de pago.",
        "fecha de creación, vencimiento y confirmación.",
        "payload de webhook recibido, firma y resultado de validación.",
        "método usado: QR, push o link.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("UX recomendada", level=1)
    for item in [
        "Desktop: mostrar QR grande, monto exacto, banco/app esperada y timer.",
        "Mobile: preferir push o link; si solo hay QR, ofrecer copiar instrucciones o cambiar método.",
        "No pedir comprobante como mecanismo principal; puede servir solo para soporte.",
        "Mostrar estado claro: esperando pago, pago recibido, vencido o hubo un problema.",
        "Mantener tarjeta/Mercado Pago como alternativa si el cliente no puede pagar con CoDi.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("Casos borde que hay que manejar", level=1)
    for item in [
        "Webhook duplicado: debe ser idempotente y no liberar dos veces.",
        "Pago por monto diferente: mandar a revisión manual.",
        "Pago después del vencimiento: decidir si aceptar, rechazar o revisar.",
        "Cliente cierra la página: el webhook debe completar la orden igual.",
        "Proveedor caído: permitir reintento o método alternativo.",
        "Reembolso: definir proceso operativo, porque puede no ser tan automático como tarjeta.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("Checklist antes de contratar proveedor", level=1)
    for item in [
        "Confirmar si la comisión por pago CoDi es $0 real.",
        "Confirmar alta, mensualidad, mínimo mensual, costo de API, CIE, portal o soporte.",
        "Confirmar sandbox y documentación técnica.",
        "Confirmar webhook de pago confirmado y método de firma/seguridad.",
        "Confirmar QR dinámico, push y/o link para mobile.",
        "Confirmar límite por operación y tiempo de liquidación.",
        "Confirmar requisitos fiscales y bancarios para México.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("Implementación por etapas", level=1)
    for item in [
        "Etapa 1: integrar sandbox, crear QR por orden y recibir webhook.",
        "Etapa 2: agregar pantalla de espera con polling y manejo de vencimiento.",
        "Etapa 3: validar casos borde, idempotencia y revisión manual.",
        "Etapa 4: liberar producción con montos chicos y monitoreo.",
        "Etapa 5: optimizar mobile con push/link si el proveedor lo soporta.",
    ]:
        add_bullet(doc, item)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build()
